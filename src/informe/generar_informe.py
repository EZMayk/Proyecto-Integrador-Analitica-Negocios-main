"""Genera el informe ejecutivo DOCX de SafeAnalytics EC."""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

RAIZ = Path(__file__).resolve().parents[2]
DATOS = RAIZ / "data" / "processed" / "homicidios_procesados.csv"
RESUMEN = RAIZ / "reports" / "resumen_analitico.json"
RANKING = RAIZ / "reports" / "ranking_riesgo_cantones.csv"
SALIDA = RAIZ / "informe" / "Informe_SafeAnalytics_EC.docx"
AZUL = RGBColor(0x15, 0x65, 0xC0)


def main() -> None:
    datos = pd.read_csv(DATOS)
    resumen = json.loads(RESUMEN.read_text(encoding="utf-8"))
    ranking = pd.read_csv(RANKING)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    portada = doc.add_paragraph()
    portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo = portada.add_run("\n\nSafeAnalytics EC\n")
    titulo.bold = True
    titulo.font.size = Pt(24)
    titulo.font.color.rgb = AZUL
    portada.add_run(
        "\nSistema Inteligente de Analítica de Negocios para el Monitoreo y "
        "Análisis Estratégico de Homicidios Intencionales en Ecuador\n\n"
        "Informe ejecutivo · Enero–mayo de 2026"
    )
    doc.add_page_break()

    def seccion(nombre: str, texto: str) -> None:
        encabezado = doc.add_heading(nombre, level=1)
        for run in encabezado.runs:
            run.font.color.rgb = AZUL
        doc.add_paragraph(texto)

    seccion(
        "Resumen ejecutivo",
        f"SafeAnalytics EC analiza {resumen['total_homicidios']:,} registros históricos "
        "de homicidios intencionales en Ecuador. Integra indicadores territoriales, "
        "temporales, delictivos y demográficos para apoyar decisiones basadas en datos.",
    )
    seccion(
        "Problemática y objetivo",
        "La dispersión de registros tabulares dificulta identificar patrones y priorizar "
        "recursos. El sistema consolida el dataset en un dashboard ejecutivo interactivo "
        "con analítica descriptiva, diagnóstica, predictiva y prescriptiva.",
    )

    doc.add_heading("KPIs principales", level=1)
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Light Shading Accent 1"
    tabla.rows[0].cells[0].text = "Indicador"
    tabla.rows[0].cells[1].text = "Resultado"
    kpis = [
        ("Total de homicidios", f"{resumen['total_homicidios']:,}"),
        ("Provincia con mayor incidencia", resumen["provincia_mayor_incidencia"][0]["provincia"]),
        ("Cantón con mayor incidencia", resumen["canton_mayor_incidencia"][0]["canton"]),
        ("Arma más registrada", resumen["arma_mas_utilizada"][0]["arma"]),
        ("Franja más frecuente", resumen["franja_mas_frecuente"][0]["franja_horaria"]),
        ("Porcentaje con arma de fuego", f"{resumen['porcentaje_arma_fuego']:.1f}%"),
    ]
    for indicador, valor in kpis:
        celdas = tabla.add_row().cells
        celdas[0].text, celdas[1].text = indicador, str(valor)

    seccion(
        "Hallazgos descriptivos y diagnósticos",
        f"La mayor incidencia se concentra en {kpis[1][1]} y el cantón principal es "
        f"{kpis[2][1]}. El arma más registrada es {kpis[3][1]} y la franja de mayor "
        f"frecuencia es {kpis[4][1]}. El análisis conjunto permite focalizar territorios, "
        "horarios, lugares y factores delictivos relevantes.",
    )
    seccion(
        "Proyección referencial",
        f"Una tendencia lineal simple estima {resumen['proyeccion_siguiente_periodo']:,} "
        "casos para el siguiente periodo si el patrón observado continuara. Esta predicción "
        "es referencial, académica y no representa un modelo criminológico profesional.",
    )

    doc.add_heading("Recomendaciones prescriptivas", level=1)
    for recomendacion in resumen["recomendaciones"]:
        doc.add_paragraph(recomendacion, style="List Bullet")
    doc.add_paragraph("Cantones prioritarios según score:")
    for fila in ranking.head(10).itertuples():
        doc.add_paragraph(
            f"{fila.canton} ({fila.provincia}): score {fila.score_riesgo}, "
            f"{fila.total_casos} casos.",
            style="List Bullet 2",
        )

    seccion(
        "Limitaciones",
        "El análisis utiliza un archivo histórico estático de enero a mayo de 2026. "
        "No mide tasas por población, subregistro, causalidad, estacionalidad anual ni "
        "efectos de intervenciones. Los resultados deben interpretarse con contexto "
        "institucional y territorial.",
    )
    seccion(
        "Conclusión ejecutiva",
        "SafeAnalytics EC permite convertir datos históricos de homicidios intencionales "
        "en información estratégica para apoyar la toma de decisiones en seguridad "
        "ciudadana, mediante visualización interactiva, detección de patrones, proyección "
        "referencial y priorización de recursos.",
    )
    doc.add_paragraph(f"Registros analizados: {len(datos):,}.")
    SALIDA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SALIDA)
    print(f"Informe generado: {SALIDA.relative_to(RAIZ)}")


if __name__ == "__main__":
    main()
